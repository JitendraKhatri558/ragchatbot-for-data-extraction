import os
import re
import fitz
from PIL import Image as PILImage
import pytesseract
import torch
from transformers import BlipProcessor, BlipForConditionalGeneration
from typing import List, Dict, Any
from docx import Document as DocxDocument

# Import from json_utils module
from json_utils import normalize_keys

# --- Vision Model Setup ---
vision_device = "cuda" if torch.cuda.is_available() else "cpu"
try:
    vision_processor = BlipProcessor.from_pretrained("Salesforce/blip-image-captioning-base", use_fast=True)
    vision_model = BlipForConditionalGeneration.from_pretrained("Salesforce/blip-image-captioning-base").to(vision_device)
    print(f"✅ Vision model loaded on {vision_device}")
except Exception as e:
    print(f"⚠️ Failed to load")
    vision_model = None
    vision_processor = None
def extract_clean_machine_names_from_products(products):
    """Extract only the cleaned list of machine names (top-level) from products."""
    result = extract_machine_names_from_products(products)
    return result.get("machines", [])
def describe_image(image_path: str) -> str:
    """Generate caption using BLIP."""
    if not os.path.exists(image_path) or vision_model is None:
        return "Image captioning unavailable"
    try:
        raw_image = PILImage.open(image_path).convert("RGB")
        inputs = vision_processor(raw_image, return_tensors="pt").to(vision_device)
        out = vision_model.generate(**inputs)
        caption = vision_processor.decode(out[0], skip_special_tokens=True)
        return caption
    except Exception as e:
        return f"Caption error: {str(e)}"

def extract_text_from_image(image_path: str) -> str:
    try:
        img = PILImage.open(image_path).convert("RGB")
        return pytesseract.image_to_string(img)
    except Exception as e:
        return f"OCR failed: {str(e)}"

def extract_images_from_docx(docx_path: str, image_dir: str) -> List[str]:
    doc = DocxDocument(docx_path)
    os.makedirs(image_dir, exist_ok=True)
    image_paths = []
    count = 0
    ns = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
        'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture'
    }
    for para in doc.paragraphs:
        p_element = para._element
        for drawing in p_element.findall('.//w:drawing', ns):
            for inline in drawing.findall('.//w:inline', ns):
                for graphic in inline.findall('.//a:graphic', ns):
                    for data in graphic.findall('.//a:graphicData', ns):
                        for pic in data.findall('.//pic:pic', ns):
                            for blip in pic.findall('.//a:blip', ns):
                                embed = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                                if embed:
                                    try:
                                        part = doc.part.related_parts[embed]
                                        ext = part.partname.split('.')[-1]
                                        img_path = os.path.join(image_dir, f"image_{count}.{ext}")
                                        with open(img_path, 'wb') as f:
                                            f.write(part.blob)
                                        image_paths.append(img_path)
                                        count += 1
                                    except Exception as e:
                                        print(f"❌ Failed to extract image {embed}: {e}")
    return image_paths

def extract_text_from_pdf(pdf_path: str, max_pages: int = 10) -> List[str]:
    """Extract text from PDF with OCR fallback."""
    texts = []
    try:
        doc = fitz.open(pdf_path)
        for page_num in range(min(max_pages, len(doc))):
            page = doc.load_page(page_num)
            text = page.get_text("text").strip()
            if len(text) < 50:  # Likely scanned
                pix = page.get_pixmap()
                img = PILImage.frombytes("RGB", [pix.width, pix.height], pix.samples)
                text = pytesseract.image_to_string(img)
            texts.append(text)
        doc.close()
    except Exception as e:
        print(f"❌ Error extracting text from PDF {pdf_path}: {e}")
        return []
    return texts

def extract_text_from_txt(file_path: str) -> str:
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        print(f"❌ Error reading TXT file {file_path}: {e}")
        return ""

def extract_products_from_docx(docx_path: str) -> List[Dict[str, Any]]:
    """
    Extract structured product data from a .docx file.
    - Metadata from key-value lines and tables
    - Text content
    - Tables
    - Images (saved and optionally captioned)
    """
    doc = DocxDocument(docx_path)
    products = []
    current_product = {
        "metadata": {},
        "text": "",
        "tables": [],
        "images": []
    }

    # Create image output dir
    image_dir = os.path.join("extracted_images", os.path.splitext(os.path.basename(docx_path))[0])
    os.makedirs(image_dir, exist_ok=True)
    image_count = 0

    # 1. Extract from paragraphs (key-value + text)
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        if ':' in text:
            key, *val = text.split(':', 1)
            key = key.strip()
            value = val[0].strip() if val else ""
            valid_keys = [
                "Title", "Company", "Price", "Condition", "Warranty", "Voltage", "Power", "Weight",
                "Model", "Brand", "Place of Origin", "Capacity", "Name of machine", "Product name",
                "Related machine", "Image URL", "source_url"
            ]
            if key in valid_keys:
                current_product["metadata"][key] = value
                continue

        # Add to general text
        current_product["text"] += text + "\n"

    # 2. Extract from tables
    for table in doc.tables:
        table_data = {}
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if len(cells) >= 2:
                key = cells[0]
                value = cells[1]
                if key:
                    table_data[key] = value
        if table_data:
            current_product["tables"].append(table_data)

    # 3. Extract images
    # Extract images
    rels = doc.part.rels
    for rel in rels:
        if "image" in rels[rel].target_ref:
            try:
                image_blob = rels[rel].target_part.blob
                ext = rels[rel].target_part.partname.split('.')[-1]
                img_dir = os.path.join("extracted_images", os.path.splitext(os.path.basename(docx_path))[0])
                os.makedirs(img_dir, exist_ok=True)
                img_path = os.path.join(img_dir, f"image_{len(current_product['images'])}.{ext}")
                with open(img_path, 'wb') as f:
                    f.write(image_blob)
                current_product["images"].append({
                    "path": img_path,
                    "caption": describe_image(img_path),
                    "source": os.path.basename(docx_path)
                })
            except Exception as e:
                print(f"❌ Failed to save image: {e}")

    # 4. Normalize metadata keys
    current_product["metadata"] = normalize_keys(current_product["metadata"])

    # 5. Add source
    current_product["metadata"]["source"] = os.path.basename(docx_path)

    # 6. Only add if metadata exists
    if current_product["metadata"]:
        products.append(current_product)

    return products

def chunk_text(text: str, chunk_size: int = 1024, chunk_overlap: int = 128) -> List[str]:
    if not text:
        return []
    words = text.split()
    chunks = []
    start = 0
    while start < len(words):
        end = min(start + chunk_size, len(words))
        chunks.append(" ".join(words[start:end]))
        if end >= len(words):
            break
        start = end - chunk_overlap
    return chunks


def process_documents(input_dir: str = "docs/rag") -> List[Dict[str, Any]]:
    if not os.path.exists(input_dir):
        print(f"❌ Directory does not exist: {input_dir}")
        return []
    print("Files in docs/rag:", os.listdir(input_dir))
    products_all = []

    for filename in os.listdir(input_dir):
        file_path = os.path.join(input_dir, filename)
        if not os.path.isfile(file_path):
            continue
        ext = os.path.splitext(filename)[-1].lower().strip(".")

        try:
            if ext == "docx":
                products = extract_products_from_docx(file_path)
                for p in products:
                    p["metadata"]["source"] = filename
                    full_text = p.get("text", "")
                    for table in p.get("tables", []):
                        if isinstance(table, dict):
                            table_text = "\n".join(
                                [f"{k}: {v}" for k, v in table.items() if k.strip()]
                            )
                            full_text += "\n" + table_text
                    p["text"] = full_text.strip()
                products_all.extend(products)

            elif ext == "txt":
                text_content = extract_text_from_txt(file_path)
                if text_content:
                    products_all.append({
                        "metadata": {"source": filename},
                        "text": text_content.strip(),
                        "tables": [],
                        "images": []
                    })

            elif ext == "pdf":
                texts_list = extract_text_from_pdf(file_path, max_pages=20)
                if texts_list:
                    products_all.append({
                        "metadata": {"source": filename},
                        "text": "\n".join(texts_list).strip(),
                        "tables": [],
                        "images": []
                    })

            else:
                print(f"📎 Unsupported file type: {filename} ({ext})")

        except Exception as e:
            print(f"❌ Error processing {filename}: {e}")

    print(f"✅ process_documents extracted {len(products_all)} products")
    return products_all


def clean_machine_name(name: str) -> str:
    if not name or len(name) < 4:
        return ""
    name = name.strip()
    prefixes_to_remove = ["the ", "this ", "our ", "a ", "name of machine:", "title:", "related machine:", "product name:"]
    for prefix in prefixes_to_remove:
        if name.lower().startswith(prefix):
            name = name[len(prefix):].strip()
    name = re.sub(r'^\d+\.\s*', '', name)
    if name.endswith('.') or name.endswith(' is') or name.endswith(' are'):
        return ""
    if any(phrase in name.lower() for phrase in ["applied to", "it can improve", "the excellence of", "brand:", "denmark", "suitable for", "large-scale", "production environments", "new condition"]):
        return ""
    name = re.sub(r"[\:\;\-\—\–\*]$", "", name).strip()
    if len(name.split()) > 15:
        return ""
    machine_keywords = ["machine", "line", "system", "conveyor", "tank", "refiner", "mixer", "pump", "tunnel", "depositor", "melter", "conche", "grinder", "coater", "dryer", "cooler", "cutter", "shear", "laser", "assembly", "packaging", "storage", "ball mill"]
    if not any(kw in name.lower() for kw in machine_keywords):
        return ""
    return name if len(name) >= 4 else ""

def extract_sub_components(full_name: str) -> List[str]:
    sub_components = []
    parts = re.split(r'[,\-\–\—\s]+', full_name)
    for part in parts:
        part = part.strip()
        if len(part) < 2:
            continue
        part_lower = part.lower()
        keywords = ["motor", "pump", "valve", "sensor", "panel", "heater", "tank", "pipe"]
        if any(kw in part_lower for kw in keywords):
            sub_components.append(part)
        elif re.search(r'\d+[kKgGLl]', part):
            sub_components.append(part)
    return sub_components

def extract_machine_names_from_products(products: List[Dict[str, Any]]) -> Dict[str, List[str]]:
    machine_keywords = ["machine", "line", "system", "conveyor", "tank", "refiner", "mixer", "pump", "tunnel", "depositor", "melter", "conche", "grinder", "coater", "dryer", "cooler", "cutter", "shear", "laser", "assembly", "packaging", "storage", "ball mill"]
    exclude_phrases = ["click here", "contact us", "terms of service", "privacy policy"]
    exclude_patterns = [r"temperature.*\(°C\)", r"rotate speed.*\(r/min\)", r"capacity.*\(kg/h\)", r"power\s*\(kw\)", r"dimension", r"voltage", r"model", r"id", r"^\d+$"]
    priority_metadata_keys = ["Name of machine", "Product name", "Related machine", "Title"]
    machine_names = set()
    sub_names = set()

    if not products:
        return {"machines": [], "sub_names": []}

    for product in products:
        if not isinstance(product, dict):
            continue
        metadata = product.get("metadata") or {}
        raw_text = product.get("text", "") or ""
        tables = product.get("tables") or []

        for key in priority_metadata_keys:
            if key in metadata:
                name = str(metadata[key]).strip()
                cleaned_name = clean_machine_name(name)
                if cleaned_name:
                    machine_names.add(cleaned_name)
                    sub_names.update(extract_sub_components(cleaned_name))

        for line in raw_text.splitlines():
            line = line.strip()
            line_lower = line.lower()
            if (any(line_lower.startswith(prefix) for prefix in ["keywords:", "processing:", "name of machine:", "related machine:", "title:", "price:", "company:", "identity:", "country:", "q1:", "a1:"]) or
                ":" in line and len(line.split(":")[0].split()) < 4 and len(line.split()) < 8 or
                re.match(r"^(q\d+:|a\d+:)", line_lower) or
                len(line) == 0 or len(line.split()) > 20 or
                any(ex in line_lower for ex in exclude_phrases) or
                any(re.search(pat, line_lower) for pat in exclude_patterns)):
                continue
            if any(kw in line_lower for kw in machine_keywords):
                cleaned_line = clean_machine_name(line)
                if cleaned_line:
                    machine_names.add(cleaned_line)
                    sub_names.update(extract_sub_components(cleaned_line))

        for table in tables:
            if isinstance(table, dict):
                name_column_key = None
                potential_name_headers = ['name', 'item', 'machine', 'component', 'product']
                for header_key in table.keys():
                    if any(potential_header in header_key.lower() for potential_header in potential_name_headers):
                        name_column_key = header_key
                        break
                if name_column_key and name_column_key in table:
                    for key, row_dict in table.items():
                        if isinstance(row_dict, dict) and name_column_key in row_dict:
                            cell_value = str(row_dict[name_column_key]).strip()
                            cleaned_value = clean_machine_name(cell_value)
                            if cleaned_value:
                                machine_names.add(cleaned_value)
                                sub_names.update(extract_sub_components(cleaned_value))
                for cell_value in table.values():
                    cell_str = str(cell_value).strip()
                    if 3 <= len(cell_str) <= 50 and len(cell_str.split()) <= 5:
                        cell_lower = cell_str.lower()
                        if any(keyword in cell_lower for keyword in ["motor", "pump", "valve", "sensor", "panel", "heater", "tank"]):
                            sub_names.add(cell_str)

    cleaned_machines = sorted({m for m in machine_names if len(m) >= 4})
    cleaned_sub_names = sorted({s for s in sub_names if len(s) >= 2})
    return {"machines": cleaned_machines, "sub_names": cleaned_sub_names}


def extract_machine_hierarchy_from_products(products):
    """
    Extracts machine hierarchy: main machine (from title) and sub-machines (from table or metadata).
    Returns grouped list and flat consolidated list.
    """
    grouped_machines = []
    all_machines = {}

    exclude_patterns = [
        r"temperature.*\(°C\)",
        r"rotate speed.*\(r/min\)",
        r"capacity.*\(kg/h\)",
        r"power.*\(kw\)",
        r"voltage.*\(v\)",
        r"weight.*\(kg\)"
    ]

    import re

    for product in products:
        metadata = product.get("metadata", {})
        text = product.get("text", "")

        # --- Step 1: Infer Main Machine from Title ---
        title = metadata.get("Title", "").strip()
        if not title:
            title = metadata.get("name_of_machine", "") or metadata.get("product_name", "")
        title = title.strip()

        if not title:
            continue  # Skip if no title

        # Assume the product title is the main machine if it contains keywords
        is_main_line = any(kw in title.lower() for kw in ["production line", "making machine", "system", "complete"])
        main_machine = title if is_main_line else f"Standalone System: {title}"

        # --- Step 2: Extract Sub-Machines ---
        sub_machines = []

        # A. From metadata fields
        for key in ["name_of_machine", "product_name", "related_machine"]:
            val = metadata.get(key, "")
            if val and val != title:
                sub_machines.append(val)

        lines = [line.strip() for line in text.split('\n') if line.strip()]
        for line in lines:
            # Match pattern: alphanumeric code followed by text
            match = re.match(r'^([A-Z]{2,}[0-9]{2,3}[A-Z0-9]*)\s+(.+)', line)
            if match:
                model, name = match.groups()
                full_name = f"{model} {name}".strip()
                if full_name not in sub_machines and full_name.lower() not in title.lower():
                    sub_machines.append(full_name)

        # C. From products_listed (if available)
        if "products_listed" in product and isinstance(product["products_listed"], list):
            for item in product["products_listed"]:
                if item not in sub_machines and item.lower() not in title.lower():
                    sub_machines.append(item)

        # --- Step 3: Clean and Deduplicate ---
        def clean_machine_name(name):
            """Clean and normalize machine name."""
            if not name or not isinstance(name, str):
                return ""
            import re
            name = re.sub(r'\s+', ' ', name.strip())
            name = re.sub(r'[^\w\s\-:/]', '', name)  # Remove bad chars
            return name if len(name) > 2 else ""

        cleaned_main = clean_machine_name(main_machine)
        cleaned_subs = [clean_machine_name(m) for m in sub_machines]
        cleaned_subs = [m for m in cleaned_subs if m and m != cleaned_main]

        # Deduplicate
        cleaned_subs = list(dict.fromkeys(cleaned_subs))

        # --- Step 4: Build Group ---
        if cleaned_main or cleaned_subs:
            group = {
                "main": [cleaned_main] if cleaned_main else [],
                "sub": cleaned_subs,
                "source": metadata.get("source", "unknown")
            }
            grouped_machines.append(group)

            # Add to global list
            if cleaned_main:
                all_machines[cleaned_main] = True
            for m in cleaned_subs:
                all_machines[m] = True

    consolidated_list = sorted(all_machines.keys())
    return grouped_machines, consolidated_list
# In json_utils.py
def format_machine_hierarchy_report(grouped: list, consolidated: list, doc_name: str = "Document") -> str:
    report = f"\n📊 Machine Hierarchy Report — {doc_name}\n"
    report += "=" * 60 + "\n\n"

    if grouped:
        report += "🔧 Main Machines & Sub-Components:\n"
        for i, group in enumerate(grouped, 1):
            mains = group.get("main", [])
            subs = group.get("sub", [])
            source = group.get("source", "unknown")

            if mains:
                report += f"{i}. 🏭 Main: {', '.join(mains)}\n"
                if subs:
                    report += f"   ⚙️ Sub: {', '.join(subs)}\n"
                report += f"   📁 Source: {source}\n"
    else:
        report += "No machines found.\n"

    report += "\n📋 All Unique Machines Found:\n"
    if consolidated:
        for m in consolidated:
            report += f" - {m}\n"
    else:
        report += " (None)\n"

    return report

